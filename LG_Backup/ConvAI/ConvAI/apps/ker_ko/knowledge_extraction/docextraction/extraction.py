# -*- coding: utf-8 -*-
"""
 * Copyright (c) 2020 LG Electronics Inc.
 * SPDX-License-Identifier: LicenseRef-LGE-Proprietary
 *
 * This program or software including the accompanying associated documentation
 * (“Software”) is the proprietary software of LG Electronics Inc. and or its
 * licensors, and may only be used, duplicated, modified or distributed pursuant
 * to the terms and conditions of a separate written license agreement between you
 * and LG Electronics Inc. (“Authorized License”). Except as set forth in an
 * Authorized License, LG Electronics Inc. grants no license (express or implied),
 * rights to use, or waiver of any kind with respect to the Software, and LG
 * Electronics Inc. expressly reserves all rights in and to the Software and all
 * intellectual property therein. If you have no Authorized License, then you have
 * no rights to use the Software in any ways, and should immediately notify LG
 * Electronics Inc. and discontinue all use of the Software.

@author: senthil.sk@lge.com
"""

from lxml import etree as et
from string import digits
from docx import Document
import sys
import os
import logging as logger
import zipfile
# TODO check relative imports are possible
#sys.path.append(os.path.abspath(os.path.join(os.path.dirname(os.path.realpath(__file__)), '..')))
from ..constants import params as cs

FONT_SIZE = "font_size"
ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
ns_pfx = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'


class StyleInfoExtractor:
    """
       This class used to parse the styles.xml and kept all the
       style info as dict
    """
    style_dict = {}

    def _read_style_info(self, etree):
        """
           read style info from the xml and converted into dict

           @param - etree object of xml
        """
        style_tags = etree.findall(".//w:style", ns)

        for style_tag in style_tags:
            styleid = style_tag.get(ns_pfx + "styleId")
            self.style_dict[styleid] = {}
            run_prop = style_tag.find(".//w:rPr", ns)
            if run_prop != None:
                size_tag = run_prop.find(".//w:sz", ns)
                if size_tag != None:
                    size_value = size_tag.get(ns_pfx + "val")
                    self.style_dict[styleid][FONT_SIZE] = size_value

    def pares_style_info(self, file_path):
        """
           parse the syle info from the styles.xml

           @param - file_path - styles.xml path
        """
        tree = et.parse(file_path)
        self._read_style_info(tree)

    def get_style_info(self, styleid):
        """
           get the style info based on the style id

           @param - styleid - styleid to fetch info
           @return - style info dict
        """
        return self.style_dict[styleid]


class SectionExtractor:
    """
       Used to extract the section wise details from the
       document.xml and maintain the details as dict
    """

    def __init__(self):
        self.section_list = list()
        self.section_dict = {}
        self.section_detail = {}
        self.info_detail = {}
        self.toc_fnd_flag = False
        self.toc_end_flag = False
        self.last_key = ""
        self.toc_font_size = 0
        self.cur_para_font_size = 0
        self.pre_para_font_size = 0
        self.style_extract = StyleInfoExtractor()

    def _check_start_with_num_rem(self, text):
        """
           check if text starts with digit and remove
           the starting digits from text

           @text - text read from docx
           @return - text without leading numerics
        """
        text = text.strip()
        flag = False
        if text[0].isdigit():
            flag = True
            text = text.lstrip(digits)
        return text, flag

    def _check_present_in_toc_sections(self, text):
        """
           Check whether the given text is present in section headers

           @text - section header text

           @return true - if the given text is section header
           @return false - if the given text is not section header
        """
        keys = self.section_dict.keys()
        text, flag = self._check_start_with_num_rem(text)
        for key in keys:
            if key == text:
                return True
        return False

    def _check_present_in_toc_sub_sections(self, text):
        """
           Check whether the given text is present in given sections sub-section

           @section - section header text
           @text - sub-section header text

           @return true - if the given text is sub-section header
           @return false - if the given text is not sub-section header
        """
        keys = self.section_dict.keys()
        text, flag = self._check_start_with_num_rem(text)

        for key in keys:
            for sub_section in self.section_dict[key]:
                # logger.debug('checking sec :  (%s),(%s),(%s)',sub_section,text,(sub_section.lower().strip() == text.lower().strip()))
                if sub_section.lower().strip() == text.lower().strip():
                    return True
        return False

    def _form_text_from_para(self, w_t_childs):
        """
           Form the text from all the <w:t> inside <w:p>

           @param - w_t_childs - List of <w:t> tags
           @return - text formed from the list of <w:t> tag
        """
        text = ""
        for w_t_child in w_t_childs:
            text = text + w_t_child.text
        return text

    def _is_toc_header(self, text):
        """
         Used to identy whether the given text is
         table of content header (English or spanish) or not

         @param - text - text from the docx
         @return true - if text param is TOC header

        """
        if text.strip() == "TABLE OF CONTENTS":
            return True
        elif text.strip() == "ÍNDICE":
            return True
        else:
            return False

    def _remove_unwanted_char(self, text):
        if text[-1] == '-':
            return text[:-1]
        else:
            return text

    def _add_section(self, text):

        """
           Add the section and sub-section from
           Table of Content in a dict.

           @param - text - text from the docx file
        """
        text, flag = self._check_start_with_num_rem(text)
        # global self.last_key

        if self.toc_end_flag == False:
            if flag:  # check string starts with number
                if text.isupper():  # checks if all char is in Upper case
                    self.last_key = text.strip()
                    if text not in self.section_dict:
                        self.section_dict[text.strip()] = []
                else:
                    if (len(self.last_key) != 0) and (self.last_key in self.section_dict):
                        self.section_list = self.section_dict[self.last_key]
                        self.section_list.append(text.strip())
            else:
                if text.isupper():
                    key = self._remove_unwanted_char(self.last_key)
                    self.section_dict[key + text.strip()] = self.section_dict.pop(self.last_key)
                    self.last_key = self.last_key + text.strip()

    def _add_section_para(self, para):

        """
           Add the section and sub-section from
           Table of Content in a dict.

           @param - text - text from the docx file
        """
        w_t_childs = para.findall(".//w:t", ns)
        text = self._form_text_from_para(w_t_childs)
        text, flag = self._check_start_with_num_rem(text)
        para_info = self._fetch_paragraph_prop(para)
        # logger.debug('debug_font_size : %s %s',para_info,text)
        if (para_info is not None) and (FONT_SIZE in para_info):
            self.cur_para_font_size = int(para_info[FONT_SIZE])
        # global self.last_key

        if self.toc_end_flag == False:
            if flag:  # check string starts with number
                # logger.debug('self.cur_para_font_size : (%s) self.pre_para_font_size(%s)',self.cur_para_font_size,self.pre_para_font_size)
                if text.isupper() and (self.cur_para_font_size > self.pre_para_font_size):  # checks if all char is in Upper case
                    self.pre_para_font_size = self.cur_para_font_size
                    self.last_key = text.strip()
                    if text not in self.section_dict:
                        self.section_dict[text.strip()] = []
                else:
                    if (len(self.last_key) != 0) and (self.last_key in self.section_dict):
                        self.pre_para_font_size = self.cur_para_font_size
                        self.section_list = self.section_dict[self.last_key]
                        self.section_list.append(text.strip())
            else:
                if text.isupper():
                    key = self._remove_unwanted_char(self.last_key)
                    # future refe self.pre_para_font_size = self.cur_para_font_size
                    # future refe logger.debug('key_str : %s', self.last_key)
                    self.section_dict[key + text.strip()] = self.section_dict.pop(self.last_key)
                    self.last_key = key + text.strip()

    def _fetch_paragraph_prop(self, paragraph):
        """
           Fetch the specific paragraph property details from styles.xml

           @param - paragraph - paragraph tag from xml
           @return - para_info - paragraph style info from styles.xml
        """
        para_prop = paragraph.find(".//w:pPr", ns)
        para_style = para_prop.find(".//w:pStyle", ns)

        if para_style != None:
            style_id = para_style.get(ns_pfx + "val")
            para_info = self.style_extract.get_style_info(style_id)
            # logger.debug('style fs : %s %s',para_info,style_id)
            return para_info
        else:
            run_prop = para_prop.find(".//w:rPr", ns)
            prop_size = run_prop.find(".//w:sz", ns)
            if prop_size is not None:
                size = prop_size.get(ns_pfx + "val")
                para_info = {}
                para_info[FONT_SIZE] = int(size)
                # logger.debug('para_size : (%s)', para_info)
                return para_info
        return None

    def _validate_paragraph(self, paragraph):
        """
           Validate the paragraph font size similiar to Table Of Content(TOC) font size

           @param - paragraph - paragraph tag from xml
           @return - true - if TOC and paragraph font size is similiar
           @return - false - if TOC and paragraph font size is not similiar
        """
        para_info = self._fetch_paragraph_prop(paragraph)
        para_font_size = -1
        if (para_info != None) and (FONT_SIZE in para_info):
            para_font_size = para_info[FONT_SIZE]

        if para_font_size == self.toc_font_size:
            return True
        return False

    def _extract_section_details(self, doc_file_path, style_file_path):
        """
           -Identified list of sections and sub-sections
           -Extract the text and add it under specific section

           @param -  doc_file_path - document.xml file path
           @param - style_file_path - style.xml file path

           @return - segregated info as section and sub-section as dict
        """
        tree = et.parse(doc_file_path)
        self.style_extract.pares_style_info(style_file_path)
        key = ""
        # extract table of content and section details from docx
        for weight in tree.findall(".//w:p", ns):
            if weight.getparent().tag == ns_pfx + "body":
                w_t_childs = weight.findall(".//w:t", ns)
                text = self._form_text_from_para(w_t_childs)
                if text == 'ENGLISH':
                    logger.debug('parent : %s', weight.getparent().tag)
                if self._is_toc_header(text.strip()):
                    # logger.debug('is_toc_header :(%s)',text)
                    self.toc_fnd_flag = True
                    self.toc_end_flag = False
                    para_info = self._fetch_paragraph_prop(weight)
                    # logger.debug('para_info : (%s)',para_info)
                    if para_info is not None:
                        # logger.debug('para_info :(%s)',para_info)
                        self.toc_font_size = para_info[FONT_SIZE]
                elif (self.toc_fnd_flag == True) and (self.toc_end_flag == False):
                    if (len(text.strip()) != 0):
                        self.toc_end_flag = self._check_present_in_toc_sections(text)
                        # logger.debug('toc_end_flag1 : (%s),(%s)',self.toc_end_flag,text)
                    if (len(text.strip()) != 0) and (self.toc_end_flag == False):
                        # self._add_section(text.strip())
                        self._add_section_para(weight)

                # Add section paragraoh under the corresponding section header in dict
                if (self.toc_end_flag == True) and (len(text.strip()) != 0):
                    flag = self._check_present_in_toc_sections(text.strip())
                    v_flag = self._validate_paragraph(weight)
                    logger.debug("v_flag : %s",v_flag)
                    # if (flag == True) and (v_flag == True):
                    if (flag == True):
                        key = text.strip()
                        self.section_detail.setdefault(key, [])
                    elif key in self.section_detail:
                        self.section_detail[key].append(text)

        return self._extract_sub_section_detail()

    def _extract_sub_section_detail(self):
        """
           segregate the sub-section details from section details

           @return - segregated details as dict
        """

        section_keys = self.section_detail.keys()
        logger.debug('toc :(%s)', self.section_dict)
        logger.debug('section : (%s)', self.section_detail)
        # logger.debug('para_text : (%s)', section_keys)
        # extract sub-section details from the section details
        for key in section_keys:
            self.info_detail[key] = {}
            last_sub_section = "sub_section"
            for section_para in self.section_detail[key]:
                if len(section_para.strip()) != 0:
                    if self._check_present_in_toc_sub_sections(section_para) == True:
                        if section_para not in self.info_detail[key]:
                            last_sub_section = section_para.lower()
                            self.info_detail[key].setdefault(last_sub_section, [])
                            logger.debug('last_sub_section : %s', last_sub_section)
                    else:
                        # if no sub_section under a section initialize section with empty list
                        if last_sub_section not in self.info_detail[key]:
                            self.info_detail[key].setdefault(last_sub_section, [])
                        self.info_detail[key][last_sub_section].append(section_para)
        return self.info_detail

    def extract_section_detail(self, doc_file_path, style_file_path):
        return self._extract_section_details(doc_file_path, style_file_path)


class DocxTableExtractor:
    """
       Extract the list of tables from the docx file
       @method get_all_tables : Get the list of tables from the document
    """

    def get_all_tables(self, file_path):
        """
            Get the list of tables from the document
            @file_path : Absolute file path of the docx file
            @return : dict of {table_name : table_obj}
        """
        table_dict = {}
        doc = Document(file_path)
        table_dict = self._extract_table(doc)
        return table_dict

    def get_section_tables(self, file_path, section):
        """
            Get table for the section from the documentation
            @file_path : Absolute file path of the docx file
            @return : dict of {table_name : table_obj}
        """
        table_dict = {}
        doc = Document(file_path)
        table_dict = self._extract_table_of_section(doc, section)
        return table_dict

    def _extract_table(self, doc):
        """
            This method dictionary of all table realated to different section of the file
            @doc: actual file from table should be extracted
        """
        section = cs.SPEC_SECTION

        # TODO - Code need to be implemented
        table = {section: doc.tables[0]}
        return table

    def _extract_table_of_section(self, doc, section):
        """
            This method dictionary of table related to different section of the file
            @doc: actual file from table should be extracted
        """
        # TODO - Code need to be implemented
        # Get all the tables from the document and fetch table related to the section
        table_dict = self._extract_table(doc)

        return table_dict[section]


if __name__ == "__main__":
    logger.basicConfig(level=logger.DEBUG,
                       format="%(asctime)s.%(msecs)03d %(levelname)s: %("
                              "funcName)s() %(message)s",
                       datefmt='%Y-%m-%d,%H:%M:%S')


    section_extractor = SectionExtractor()
    with open("section_details.txt", 'w') as f:
        f.write(str(section_extractor.extract_section_detail(
            "E:\work_from_home\manuals\washing_machine\WM4500H_A\manual_washing_machine_test\word\document.xml",
            "E:\work_from_home\manuals\washing_machine\WM4500H_A\manual_washing_machine_test\word\styles.xml")))
